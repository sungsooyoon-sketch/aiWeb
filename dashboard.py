"""
장납기 자재 공급 리스크 분석 - Streamlit 대시보드
실행: streamlit run dashboard.py
"""

import os
import io
import json
from pathlib import Path
from datetime import datetime

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── 환경 설정 ───────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent
load_dotenv(BASE_DIR / ".env")

OPENAI_API_KEY = os.getenv("OPENAI_SERVICE_KEY")
client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None

FILES = {
    "bom":       BASE_DIR / "bom_data.xlsx",
    "inventory": BASE_DIR / "inventory_data.xlsx",
    "material":  BASE_DIR / "material_master.xlsx",
    "order":     BASE_DIR / "order_data.xlsx",
}

RISK_COLOR = {"심각": "#C00000", "주의": "#ED7D31", "양호": "#375623"}
RISK_BG    = {"심각": "#FCE4D6", "주의": "#FFEB9C", "양호": "#E2EFDA"}


# ─── 데이터 로드 (캐시) ──────────────────────────────────────────────────────
@st.cache_data(ttl=60)
def load_data():
    return {
        "bom":       pd.read_excel(FILES["bom"],       sheet_name="BOM"),
        "inventory": pd.read_excel(FILES["inventory"], sheet_name="현재고"),
        "material":  pd.read_excel(FILES["material"],  sheet_name="자재마스터"),
        "order":     pd.read_excel(FILES["order"],     sheet_name="수주데이터"),
    }


# ─── 장납기 자재 필요 수량 산출 ─────────────────────────────────────────────
@st.cache_data(ttl=60)
def calc_requirements(_dfs):
    orders_active = _dfs["order"][_dfs["order"]["수주상태"] == "진행중"].copy()
    order_qty = (
        orders_active.groupby("제품코드", as_index=False)["수주수량"]
        .sum()
        .rename(columns={"수주수량": "총수주수량"})
    )
    bom_order = order_qty.merge(_dfs["bom"], on="제품코드", how="inner")
    bom_order["필요수량"] = bom_order["총수주수량"] * bom_order["소요수량"]
    req = bom_order.groupby("자재코드", as_index=False)["필요수량"].sum()
    req = req.merge(_dfs["material"], on="자재코드", how="inner")
    long_lead = req[req["장납기여부"] == "Y"].copy()
    long_lead = long_lead.merge(
        _dfs["inventory"][["자재코드", "현재고수량", "안전재고수량", "재고상태"]],
        on="자재코드", how="left",
    )
    long_lead["현재고수량"]  = long_lead["현재고수량"].fillna(0)
    long_lead["안전재고수량"] = long_lead["안전재고수량"].fillna(0)
    long_lead["부족수량"] = (long_lead["필요수량"] - long_lead["현재고수량"]).clip(lower=0)
    long_lead["충족률(%)"] = (
        (long_lead["현재고수량"] / long_lead["필요수량"].replace(0, 1)) * 100
    ).clip(upper=100).round(1)

    def risk_level(row):
        if row["부족수량"] > 0:
            return "심각"
        if row["현재고수량"] <= row["안전재고수량"]:
            return "주의"
        return "양호"

    long_lead["위험등급"] = long_lead.apply(risk_level, axis=1)
    return long_lead.reset_index(drop=True)


# ─── OpenAI 분석 ─────────────────────────────────────────────────────────────
def generate_ai_report(df: pd.DataFrame) -> str:
    materials_info = []
    for _, row in df.iterrows():
        materials_info.append({
            "자재코드": row["자재코드"],
            "자재명": row["자재명"],
            "자재분류": row["자재분류"],
            "협력사명": row["협력사명"],
            "협력사국가": row["협력사국가"],
            "리드타임(일)": int(row["리드타임(일)"]),
            "비고": str(row["비고"]) if pd.notna(row.get("비고")) else "",
            "필요수량": int(row["필요수량"]),
            "현재고수량": int(row["현재고수량"]),
            "안전재고수량": int(row["안전재고수량"]),
            "부족수량": int(row["부족수량"]),
            "위험등급": row["위험등급"],
        })
    critical = [m for m in materials_info if m["위험등급"] == "심각"]
    caution  = [m for m in materials_info if m["위험등급"] == "주의"]
    normal   = [m for m in materials_info if m["위험등급"] == "양호"]

    prompt = f"""
당신은 제조업 공급망 리스크 분석 전문가입니다.
아래 장납기 자재 현황 데이터를 분석하여 경영진을 위한 공급 리스크 분석 리포트를 작성해 주세요.

## 장납기 자재 현황 데이터
{json.dumps(materials_info, ensure_ascii=False, indent=2)}

## 요약 통계
- 장납기 자재 총 {len(materials_info)}종 / 심각: {len(critical)}종 / 주의: {len(caution)}종 / 양호: {len(normal)}종

## 리포트 작성 요구사항
다음 섹션으로 구성된 전문적인 리포트를 한국어로 작성하세요:

1. **경영진 요약 (Executive Summary)**
2. **위험 자재별 상세 분석** (심각 등급 우선, 단일공급사 리스크 포함)
3. **공급망 취약점 분석** (국가별·분류별 집중도)
4. **긴급 대응 방안** (즉각 조치, 긴급 발주 권고)
5. **중장기 리스크 완화 전략**
"""
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "당신은 제조업 공급망 리스크 분석 전문가입니다."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.3,
        max_tokens=3000,
    )
    return response.choices[0].message.content


# ─── Word 문서 생성 (bytes 반환) ─────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def build_word_bytes(df: pd.DataFrame, ai_analysis: str) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.page_width    = Cm(29.7)
    section.page_height   = Cm(21.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # 제목
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("장납기 자재 공급 리스크 분석 리포트")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"작성일: {datetime.now().strftime('%Y년 %m월 %d일')}   |   기준: 수주데이터, BOM, 자재마스터, 현재고")
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    doc.add_paragraph()

    # 종합 현황 표
    h1 = doc.add_paragraph()
    h1r = h1.add_run("■ 종합 현황")
    h1r.bold = True; h1r.font.size = Pt(13)
    h1r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    stat_t = doc.add_table(rows=2, cols=4)
    stat_t.style = "Table Grid"
    stat_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdrs   = ["장납기 자재 총계", "심각 (재고 부족)", "주의 (안전재고↓)", "양호"]
    counts = [str(len(df)),
              str((df["위험등급"] == "심각").sum()),
              str((df["위험등급"] == "주의").sum()),
              str((df["위험등급"] == "양호").sum())]
    bgs    = ["2E4057", "C00000", "ED7D31", "375623"]
    for i, (h, c, bg) in enumerate(zip(hdrs, counts, bgs)):
        hc = stat_t.cell(0, i); hc.text = h
        _set_cell_bg(hc, bg)
        for run in hc.paragraphs[0].runs:
            run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        vc = stat_t.cell(1, i); vc.text = c
        for run in vc.paragraphs[0].runs:
            run.bold = True; run.font.size = Pt(18)
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 상세 현황 표
    h2 = doc.add_paragraph()
    h2r = h2.add_run("■ 장납기 자재 필요 수량 및 재고 현황")
    h2r.bold = True; h2r.font.size = Pt(13)
    h2r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    cols_display = ["자재코드", "자재명", "분류", "협력사", "국가", "리드타임", "필요수량", "현재고", "안전재고", "부족수량", "충족률(%)", "위험등급"]
    col_map      = {"자재코드": "자재코드", "자재명": "자재명", "분류": "자재분류", "협력사": "협력사명",
                    "국가": "협력사국가", "리드타임": "리드타임(일)", "필요수량": "필요수량",
                    "현재고": "현재고수량", "안전재고": "안전재고수량", "부족수량": "부족수량",
                    "충족률(%)": "충족률(%)", "위험등급": "위험등급"}
    right_cols = {"필요수량", "현재고", "안전재고", "부족수량", "리드타임", "충족률(%)"}

    det_t = doc.add_table(rows=1, cols=len(cols_display))
    det_t.style = "Table Grid"
    for i, col in enumerate(cols_display):
        c = det_t.rows[0].cells[i]; c.text = col
        _set_cell_bg(c, "1F497D")
        for run in c.paragraphs[0].runs:
            run.bold = True; run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    risk_bg = {"심각": "FCE4D6", "주의": "FFEB9C", "양호": "E2EFDA"}
    for _, row in df.iterrows():
        cells = det_t.add_row().cells
        bg = risk_bg.get(row["위험등급"], "FFFFFF")
        for i, cd in enumerate(cols_display):
            val = row[col_map[cd]]
            val = "" if pd.isna(val) else (int(val) if isinstance(val, float) else val)
            cells[i].text = str(val)
            _set_cell_bg(cells[i], bg)
            para = cells[i].paragraphs[0]
            for run in para.runs:
                run.font.size = Pt(8.5)
            if cd in right_cols:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()

    # AI 분석
    h3 = doc.add_paragraph()
    h3r = h3.add_run("■ AI 공급 리스크 분석")
    h3r.bold = True; h3r.font.size = Pt(13)
    h3r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    for line in ai_analysis.split("\n"):
        para = doc.add_paragraph()
        line = line.strip()
        if line.startswith("## ") or line.startswith("# "):
            run = para.add_run(line.lstrip("#").strip())
            run.bold = True; run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif line.startswith("**") and line.endswith("**"):
            run = para.add_run(line.strip("*"))
            run.bold = True; run.font.size = Pt(11)
        elif line.startswith("- ") or line.startswith("• "):
            para.style = "List Bullet"
            parts = line[2:].split("**")
            for idx, part in enumerate(parts):
                if part:
                    run = para.add_run(part)
                    run.bold = (idx % 2 == 1)
                    run.font.size = Pt(10)
        elif line:
            parts = line.split("**")
            for idx, part in enumerate(parts):
                if part:
                    run = para.add_run(part)
                    run.bold = (idx % 2 == 1)
                    run.font.size = Pt(10)
        else:
            para.add_run("").font.size = Pt(6)

    # 면책 문구
    doc.add_paragraph()
    note = doc.add_paragraph()
    nr = note.add_run("※ 본 리포트는 OpenAI gpt-4o-mini 모델을 활용하여 자동 생성된 문서입니다. 최종 의사결정 전 담당자의 검토를 권장합니다.")
    nr.font.size = Pt(8); nr.italic = True
    nr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── Streamlit 앱 ─────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="장납기 자재 공급 리스크 대시보드",
    page_icon="📦",
    layout="wide",
    initial_sidebar_state="expanded",
)

# CSS
st.markdown("""
<style>
    .main-header { font-size: 1.8rem; font-weight: 700; color: #1F497D; margin-bottom: 0.2rem; }
    .sub-header  { font-size: 0.9rem; color: #666; margin-bottom: 1.5rem; }
    .metric-card { background: #f8f9fa; border-radius: 10px; padding: 1rem; border-left: 4px solid; text-align: center; }
    .risk-critical { border-color: #C00000; }
    .risk-caution  { border-color: #ED7D31; }
    .risk-good     { border-color: #375623; }
    .risk-total    { border-color: #1F497D; }
    .section-title { font-size: 1.1rem; font-weight: 600; color: #1F497D;
                     border-bottom: 2px solid #1F497D; padding-bottom: 0.3rem; margin: 1rem 0 0.8rem; }
    div[data-testid="stSidebarContent"] { background: #f0f4fa; }
</style>
""", unsafe_allow_html=True)


# ─── 사이드바 ─────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("## 📦 메뉴")
    menu = st.radio(
        "페이지 선택",
        ["📊 대시보드", "📋 자재 상세", "📈 수주 현황", "📄 Word 리포트 생성"],
        label_visibility="collapsed",
    )
    st.markdown("---")
    st.markdown("**데이터 새로고침**")
    if st.button("🔄 데이터 재로드", width="stretch"):
        st.cache_data.clear()
        st.rerun()
    st.markdown("---")
    st.caption(f"마지막 로드: {datetime.now().strftime('%H:%M:%S')}")
    if not OPENAI_API_KEY:
        st.error("OPENAI_SERVICE_KEY 미설정")


# ─── 데이터 로드 ─────────────────────────────────────────────────────────────
try:
    dfs = load_data()
    df  = calc_requirements(dfs)
except Exception as e:
    st.error(f"데이터 로드 실패: {e}")
    st.stop()


# ══════════════════════════════════════════════════════════════════════════════
# 페이지 1: 대시보드
# ══════════════════════════════════════════════════════════════════════════════
if menu == "📊 대시보드":
    st.markdown('<div class="main-header">📦 장납기 자재 공급 리스크 대시보드</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="sub-header">기준일: {datetime.now().strftime("%Y년 %m월 %d일")} &nbsp;|&nbsp; 진행중 수주 기반 산출</div>', unsafe_allow_html=True)

    # KPI 카드
    n_total    = len(df)
    n_critical = (df["위험등급"] == "심각").sum()
    n_caution  = (df["위험등급"] == "주의").sum()
    n_good     = (df["위험등급"] == "양호").sum()
    total_short = int(df["부족수량"].sum())

    c1, c2, c3, c4, c5 = st.columns(5)
    with c1:
        st.markdown(f'<div class="metric-card risk-total"><div style="font-size:2rem;font-weight:700;color:#1F497D">{n_total}</div><div style="color:#555">장납기 자재 총계</div></div>', unsafe_allow_html=True)
    with c2:
        st.markdown(f'<div class="metric-card risk-critical"><div style="font-size:2rem;font-weight:700;color:#C00000">{n_critical}</div><div style="color:#555">🔴 심각 (재고 부족)</div></div>', unsafe_allow_html=True)
    with c3:
        st.markdown(f'<div class="metric-card risk-caution"><div style="font-size:2rem;font-weight:700;color:#ED7D31">{n_caution}</div><div style="color:#555">🟡 주의 (안전재고↓)</div></div>', unsafe_allow_html=True)
    with c4:
        st.markdown(f'<div class="metric-card risk-good"><div style="font-size:2rem;font-weight:700;color:#375623">{n_good}</div><div style="color:#555">🟢 양호</div></div>', unsafe_allow_html=True)
    with c5:
        st.markdown(f'<div class="metric-card risk-critical"><div style="font-size:2rem;font-weight:700;color:#C00000">{total_short:,}</div><div style="color:#555">총 부족 수량</div></div>', unsafe_allow_html=True)

    st.markdown("")

    # 차트 행 1
    col_l, col_r = st.columns([3, 2])

    with col_l:
        st.markdown('<div class="section-title">자재별 필요수량 vs 현재고수량</div>', unsafe_allow_html=True)
        chart_df = df.sort_values("부족수량", ascending=False)
        fig_bar = go.Figure()
        fig_bar.add_trace(go.Bar(name="현재고수량", x=chart_df["자재명"], y=chart_df["현재고수량"],
                                  marker_color="#4472C4"))
        fig_bar.add_trace(go.Bar(name="필요수량", x=chart_df["자재명"], y=chart_df["필요수량"],
                                  marker_color="#ED7D31", opacity=0.7))
        fig_bar.update_layout(barmode="overlay", height=320, margin=dict(t=10, b=60),
                               legend=dict(orientation="h", y=1.05),
                               xaxis_tickangle=-30)
        st.plotly_chart(fig_bar, width='stretch')

    with col_r:
        st.markdown('<div class="section-title">위험 등급 분포</div>', unsafe_allow_html=True)
        risk_counts = df["위험등급"].value_counts().reset_index()
        risk_counts.columns = ["위험등급", "count"]
        color_map = {"심각": "#C00000", "주의": "#ED7D31", "양호": "#375623"}
        fig_pie = px.pie(risk_counts, names="위험등급", values="count",
                          color="위험등급", color_discrete_map=color_map,
                          hole=0.45)
        fig_pie.update_traces(textposition="inside", textinfo="percent+label")
        fig_pie.update_layout(height=320, margin=dict(t=10, b=10),
                               showlegend=False)
        st.plotly_chart(fig_pie, width='stretch')

    # 차트 행 2
    col_l2, col_r2 = st.columns([2, 3])

    with col_l2:
        st.markdown('<div class="section-title">국가별 장납기 자재 현황</div>', unsafe_allow_html=True)
        country_df = df.groupby("협력사국가").agg(
            자재수=("자재코드", "count"),
            부족수량합=("부족수량", "sum"),
        ).reset_index()
        fig_country = px.bar(country_df, x="협력사국가", y="자재수",
                              color="부족수량합", color_continuous_scale="Reds",
                              labels={"부족수량합": "총 부족수량"})
        fig_country.update_layout(height=280, margin=dict(t=10, b=10),
                                   coloraxis_showscale=False)
        st.plotly_chart(fig_country, width='stretch')

    with col_r2:
        st.markdown('<div class="section-title">자재별 재고 충족률 (%)</div>', unsafe_allow_html=True)
        gauge_df = df.sort_values("충족률(%)")[["자재명", "충족률(%)", "위험등급"]]
        colors = [RISK_COLOR.get(r, "#999") for r in gauge_df["위험등급"]]
        fig_gauge = go.Figure(go.Bar(
            x=gauge_df["충족률(%)"], y=gauge_df["자재명"],
            orientation="h", marker_color=colors,
            text=gauge_df["충족률(%)"].astype(str) + "%",
            textposition="outside",
        ))
        fig_gauge.add_vline(x=100, line_dash="dash", line_color="gray")
        fig_gauge.update_layout(height=280, margin=dict(t=10, b=10, r=60),
                                 xaxis=dict(range=[0, 120]))
        st.plotly_chart(fig_gauge, width='stretch')

    # 심각 자재 요약 테이블
    st.markdown('<div class="section-title">🔴 심각 자재 요약</div>', unsafe_allow_html=True)
    critical_df = df[df["위험등급"] == "심각"][
        ["자재코드", "자재명", "협력사명", "협력사국가", "리드타임(일)", "필요수량", "현재고수량", "부족수량", "충족률(%)"]
    ].sort_values("부족수량", ascending=False)
    st.dataframe(
        critical_df.style.background_gradient(subset=["부족수량"], cmap="Reds")
                         .format({"충족률(%)": "{:.1f}%", "필요수량": "{:,}", "현재고수량": "{:,}", "부족수량": "{:,}"}),
        width='stretch', hide_index=True,
    )


# ══════════════════════════════════════════════════════════════════════════════
# 페이지 2: 자재 상세
# ══════════════════════════════════════════════════════════════════════════════
elif menu == "📋 자재 상세":
    st.markdown('<div class="main-header">📋 장납기 자재 상세 현황</div>', unsafe_allow_html=True)

    # 필터
    col_f1, col_f2, col_f3 = st.columns(3)
    with col_f1:
        risk_filter = st.multiselect("위험등급 필터", ["심각", "주의", "양호"], default=["심각", "주의", "양호"])
    with col_f2:
        country_filter = st.multiselect("국가 필터", df["협력사국가"].unique().tolist(),
                                         default=df["협력사국가"].unique().tolist())
    with col_f3:
        search = st.text_input("자재명 검색", "")

    filtered = df[df["위험등급"].isin(risk_filter) & df["협력사국가"].isin(country_filter)]
    if search:
        filtered = filtered[filtered["자재명"].str.contains(search, na=False)]

    st.markdown(f"**{len(filtered)}개** 자재 표시 중")

    display_cols = ["자재코드", "자재명", "자재분류", "협력사명", "협력사국가",
                    "리드타임(일)", "필요수량", "현재고수량", "안전재고수량", "부족수량", "충족률(%)", "위험등급", "재고상태"]

    def color_risk(val):
        colors = {"심각": "background-color:#FCE4D6", "주의": "background-color:#FFEB9C", "양호": "background-color:#E2EFDA"}
        return colors.get(val, "")

    styled = (
        filtered[display_cols]
        .style
        .applymap(color_risk, subset=["위험등급"])
        .format({"충족률(%)": "{:.1f}%", "필요수량": "{:,}", "현재고수량": "{:,}", "부족수량": "{:,}"})
    )
    st.dataframe(styled, width='stretch', hide_index=True, height=500)

    # 자재 개별 카드
    st.markdown('<div class="section-title">자재 상세 카드</div>', unsafe_allow_html=True)
    selected_mat = st.selectbox("자재 선택", filtered["자재명"].tolist())
    if selected_mat:
        row = filtered[filtered["자재명"] == selected_mat].iloc[0]
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("필요수량",   f"{int(row['필요수량']):,}")
        c2.metric("현재고수량", f"{int(row['현재고수량']):,}", delta=f"{int(row['현재고수량']-row['필요수량']):,}")
        c3.metric("안전재고수량", f"{int(row['안전재고수량']):,}")
        c4.metric("부족수량",   f"{int(row['부족수량']):,}", delta_color="inverse", delta=f"-{int(row['부족수량']):,}" if row['부족수량'] > 0 else "충족")

        info_cols = st.columns(2)
        with info_cols[0]:
            st.info(f"**협력사:** {row['협력사명']}  \n**국가:** {row['협력사국가']}  \n**리드타임:** {int(row['리드타임(일)'])}일  \n**분류:** {row['자재분류']}")
        with info_cols[1]:
            bom_used = dfs["bom"][dfs["bom"]["자재코드"] == row["자재코드"]]
            if not bom_used.empty:
                st.write("**이 자재를 사용하는 제품:**")
                st.dataframe(bom_used[["제품코드", "제품명", "소요수량"]], hide_index=True, width='stretch')


# ══════════════════════════════════════════════════════════════════════════════
# 페이지 3: 수주 현황
# ══════════════════════════════════════════════════════════════════════════════
elif menu == "📈 수주 현황":
    st.markdown('<div class="main-header">📈 수주 현황</div>', unsafe_allow_html=True)

    orders = dfs["order"]
    active = orders[orders["수주상태"] == "진행중"]

    kc1, kc2, kc3, kc4 = st.columns(4)
    kc1.metric("전체 수주", len(orders))
    kc2.metric("진행중 수주", len(active))
    kc3.metric("총 수주금액", f"{int(active['수주금액(원)'].sum()):,}원")
    kc4.metric("총 수주수량", f"{int(active['수주수량'].sum()):,}")

    col_a, col_b = st.columns(2)
    with col_a:
        st.markdown('<div class="section-title">제품군별 수주 수량</div>', unsafe_allow_html=True)
        group_df = active.groupby("제품군")["수주수량"].sum().reset_index()
        fig = px.bar(group_df, x="제품군", y="수주수량", color="제품군")
        fig.update_layout(height=300, margin=dict(t=10, b=10), showlegend=False)
        st.plotly_chart(fig, width='stretch')

    with col_b:
        st.markdown('<div class="section-title">거래처별 수주금액 Top 10</div>', unsafe_allow_html=True)
        cust_df = active.groupby("거래처명")["수주금액(원)"].sum().nlargest(10).reset_index()
        fig2 = px.bar(cust_df, x="수주금액(원)", y="거래처명", orientation="h",
                       color="수주금액(원)", color_continuous_scale="Blues")
        fig2.update_layout(height=300, margin=dict(t=10, b=10), coloraxis_showscale=False)
        st.plotly_chart(fig2, width='stretch')

    st.markdown('<div class="section-title">진행중 수주 목록</div>', unsafe_allow_html=True)
    st.dataframe(
        active[["수주번호", "제품코드", "제품명", "제품군", "거래처명", "담당영업", "수주수량", "수주금액(원)", "납기요청일"]]
        .sort_values("수주금액(원)", ascending=False)
        .style.format({"수주수량": "{:,}", "수주금액(원)": "{:,}"}),
        width='stretch', hide_index=True,
    )


# ══════════════════════════════════════════════════════════════════════════════
# 페이지 4: Word 리포트 생성
# ══════════════════════════════════════════════════════════════════════════════
elif menu == "📄 Word 리포트 생성":
    st.markdown('<div class="main-header">📄 Word 리포트 생성</div>', unsafe_allow_html=True)
    st.markdown("OpenAI gpt-4o-mini 모델로 공급 리스크를 분석하고 Word 문서를 생성합니다.")

    # 현황 미리보기
    st.markdown('<div class="section-title">분석 대상 자재 현황</div>', unsafe_allow_html=True)
    preview_cols = ["자재코드", "자재명", "협력사명", "협력사국가", "리드타임(일)",
                    "필요수량", "현재고수량", "부족수량", "충족률(%)", "위험등급"]

    def color_risk(val):
        colors = {"심각": "background-color:#FCE4D6;color:#C00000;font-weight:bold",
                  "주의": "background-color:#FFEB9C;color:#7F4F00",
                  "양호": "background-color:#E2EFDA;color:#375623"}
        return colors.get(val, "")

    st.dataframe(
        df[preview_cols].style
            .applymap(color_risk, subset=["위험등급"])
            .format({"충족률(%)": "{:.1f}%", "필요수량": "{:,}", "현재고수량": "{:,}", "부족수량": "{:,}"}),
        width='stretch', hide_index=True,
    )

    st.markdown("---")

    col_btn, col_info = st.columns([1, 3])
    with col_btn:
        generate_btn = st.button("🤖 AI 분석 + Word 생성", type="primary", width='stretch')
    with col_info:
        st.info("AI 분석에 약 20~40초 소요됩니다. 생성 후 다운로드 버튼이 나타납니다.")

    if generate_btn:
        if not OPENAI_API_KEY:
            st.error(".env 파일에 OPENAI_SERVICE_KEY가 없습니다.")
        else:
            with st.status("분석 진행 중...", expanded=True) as status:
                st.write("📊 데이터 준비 중...")
                st.write(f"  - 장납기 자재 {len(df)}종, 심각 {(df['위험등급']=='심각').sum()}종")

                st.write("🤖 OpenAI API 호출 중 (gpt-4o-mini)...")
                try:
                    ai_text = generate_ai_report(df)
                    st.write("  - AI 분석 완료")

                    st.write("📝 Word 문서 작성 중...")
                    word_bytes = build_word_bytes(df, ai_text)
                    fname = f"공급리스크분석리포트_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

                    st.session_state["word_bytes"] = word_bytes
                    st.session_state["word_fname"] = fname
                    st.session_state["ai_text"]    = ai_text
                    status.update(label="완료!", state="complete")
                except Exception as e:
                    status.update(label="오류 발생", state="error")
                    st.error(f"오류: {e}")

    # 결과 표시
    if "word_bytes" in st.session_state:
        st.success(f"Word 문서가 준비되었습니다: {st.session_state['word_fname']}")
        st.download_button(
            label="⬇️  Word 문서 다운로드",
            data=st.session_state["word_bytes"],
            file_name=st.session_state["word_fname"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            width='content',
        )

        st.markdown('<div class="section-title">AI 분석 결과 미리보기</div>', unsafe_allow_html=True)
        with st.expander("전체 분석 내용 보기", expanded=True):
            st.markdown(st.session_state["ai_text"])
